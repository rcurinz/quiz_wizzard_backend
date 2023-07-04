# pip install protobuf      ->    posiblemente instalar version 3.20.x o menor
# pip install pandas
import json
import os
import math
import random
import shutil
import uuid
from io import BytesIO
import time
import pandas as pd
from PyPDF2 import PdfReader
import fitz
import secrets
from ast import literal_eval
from datetime import datetime, timedelta
from flask import Flask, request, jsonify, current_app, send_from_directory
from flask import current_app, send_file
from werkzeug.security import generate_password_hash as genph
from werkzeug.security import check_password_hash as checkph
from models import *
from parrot import Parrot
import warnings
from werkzeug.utils import secure_filename
from deep_translator import GoogleTranslator
from docx import Document
from fpdf import FPDF
# Transformers
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, AutoModelWithLMHead
import torch

warnings.filterwarnings("ignore")
from transformers import AutoModelForQuestionAnswering, AutoTokenizer, pipeline
from docx2pdf import convert
import aspose.words as aw

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Flowable
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# To load the model and tokenizer
modelA = AutoModelForQuestionAnswering.from_pretrained("MarcBrun/ixambert-finetuned-squad-eu-en")
tokenizerA = AutoTokenizer.from_pretrained("MarcBrun/ixambert-finetuned-squad-eu-en")

tokenizerQ = AutoTokenizer.from_pretrained("b2bFiles")
modelQ = AutoModelForSeq2SeqLM.from_pretrained("b2bFiles")

# To save it once you download it
# tokenizerQ = AutoTokenizer.from_pretrained("mrm8488/bert2bert-spanish-question-generation")
# modelQ = AutoModelWithLMHead.from_pretrained("mrm8488/bert2bert-spanish-question-generation")

# tokenizerQ.save_pretrained("b2bFiles")
# modelQ.save_pretrained("b2bFiles")

# Cargar pipeline una vez al inicio
qa_pipeline = pipeline("question-answering", model=modelA, tokenizer=tokenizerA)

modelA = modelA#.half().to('cuda')
modelQ = modelQ#.half().to('cuda')
def prueba():
    return {"test ": "prueba"}


def generate(texto1):
    texto = texto1#literal_eval(texto1)['question']
    parrot = Parrot(model_tag="prithivida/parrot_paraphraser_on_T5")

    phrases = []
    translated = GoogleTranslator(source='spanish', target='english').translate(
        texto)
    phrases.append(translated)
    frases = []
    frases_bruto = []
    for phrase in phrases:
        print("-" * 100)
        print("Input_phrase: ", phrase)
        print("-" * 100)
        para_phrases = parrot.augment(input_phrase=phrase, use_gpu=False, do_diverse=True)
        for para_phrase in para_phrases:
            p = GoogleTranslator(source='english', target='spanish').translate(para_phrase[0])
            if p not in frases_bruto:
                frases_bruto.append(p)
                frases.append({"respuesta": p})
            print(p)
    return frases


def upload_files(files, id, id_project):
    # Obtener el usuario y proyecto de la base de datos
    user = Usuario.query.get(id)
    project = Proyecto.query.get(id_project)

    # Verificar si el usuario y proyecto existen
    if user is None or project is None:
        return [{"message": "Usuario o proyecto no encontrado"}]

    # Verificar si el proyecto pertenece al usuario
    if project.user_id != user.user_id:
        return [{"message": "El proyecto no pertenece al usuario"}]

    # Obtener un nombre de archivo seguro y generar un nombre único
    filename = secure_filename(files.filename)
    timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    new_filename = f"{id}_{id_project}_{timestamp}_{filename}"

    # Intentar guardar el archivo
    try:
        files.save(os.path.join(current_app.config['UPLOADS'], new_filename))

        # Crear un nuevo registro de archivo en la base de datos
        new_file_record = Archivo(
            project_id=id_project,
            filename=new_filename,
            original_filename=filename,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )

        # Agregar y hacer commit del nuevo registro en la base de datos
        db.session.add(new_file_record)
        db.session.commit()
    except Exception as e:
        return [{"message": f"Error al guardar archivo: {str(e)}"}]

    return [{"message": "Archivo subido correctamente", "filename": new_filename}]


def read_pdf(file, pInicio=None, pFin=None):
    reader = PdfReader(file)
    number_of_pages = len(reader.pages)

    # Si pInicio o pFin no se especifican, establecer valores predeterminados
    if pInicio is None:
        pInicio = 0
    else:
        pInicio = max(0, int(pInicio) - 1)  # Restar 1 para hacerlo base 0 y asegurarse de que no sea negativo

    if pFin is None:
        pFin = number_of_pages
    else:
        pFin = min(number_of_pages, int(pFin))  # Asegurarse de que no exceda el número de páginas

    # Limitar la cantidad de páginas a leer a 15 como máximo
    max_pages_to_read = 15
    if pFin - pInicio > max_pages_to_read:
        pFin = pInicio + max_pages_to_read

    paragraphs = []
    for k in range(pInicio, pFin):
        page = reader.pages[k]
        paragraphs.append(page.extract_text().strip('\n'))

    return paragraphs


def read_docs(file):
    doc = Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return "\n".join(fullText)


def login_user(email, password):
    user = Usuario.query.filter_by(email=email).first()

    if user is None or not checkph(user.password, password):
        return {'status': '500', "message": "Usuario o contraseña incorrectos"}

    token = secrets.token_hex(16)  # Considera encriptar este token
    user.token = token
    db.session.commit()

    hora_start = datetime.utcnow()
    hora_finish = hora_start + timedelta(hours=1)
    token_encriptado = secrets.token_hex()  # Nuevamente, considera encriptar este token

    token_registro = Token(token=token_encriptado, user_id=user.user_id, status_token='active', start_at=hora_start, finish_at=hora_finish)
    db.session.add(token_registro)
    db.session.commit()

    return {"status": "200", "message": "Usuario logeado correctamente", "token": token_encriptado, "id": user.user_id}


def register_user(data):
    username = data['name']
    email = data['email']
    password = data['password']
    hash_password = genph(password)
    user = Usuario(
        name_user=username,
        email=email,
        password=hash_password,
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow()
    )
    db.session.add(user)
    db.session.commit()
    return {"status": "200", "message": "Usuario registrado correctamente"}


def generate_answer(context: str, max_length: int = 512) -> str:
    inputText = "context: %s </s>" % (context)
    features = tokenizerQ([inputText], return_tensors='pt')

    # Mover los tensores de entrada a la GPU
    input_ids = features['input_ids']#.to('cuda')
    attention_mask = features['attention_mask']#.half().to('cuda')

    output = modelQ.generate(input_ids=input_ids,
                             attention_mask=attention_mask,
                             max_length=max_length)

    question = tokenizerQ.decode(output[0]).strip("[SEP]")
    question = question.strip("CLS]")
    #qa = pipeline("question-answering", model=modelA, tokenizer=tokenizerA)
    #pred = qa(question=question, context=context)
    pred = qa_pipeline(question=question, context=context)
    return {'status': 200,
            'message': 'ok',
            'question': question,
            'answer': pred['answer'],
            'score': pred['score']
            }


def create_new_project(data):
    name = data['name']
    description = data['description']
    user_id = data['id_user']
    project = Proyecto(
        project_name=name,
        description=description,
        user_id=user_id,
        status="active",
        created_at=datetime.now(),
        updated_at=datetime.now()
    )

    try:
        db.session.add(project)
        db.session.commit()
    except Exception as e:
        return {"status": "500", "message": str(e)}

    return {"status": "200", "message": "Proyecto creado correctamente", "project_id": project.project_id, 'name': name}


def get_projects_user(data):
    id_user = data['id_user']
    id_project = data['id_project']
    data = []
    # cuando el id del proyecto es -1 se devuelven todos los proyectos del usuario
    if id_project == -1:
        projects = Proyecto.query.filter_by(user_id=id_user).all()
        for project in projects:
            data.append({
                "id": project.project_id,
                "name": project.project_name,
                "description": project.description,
                "status": project.status,
                "created_at": project.created_at,
                "updated_at": project.updated_at,
            })
    # En caso contrario se devuelve el proyecto con el id especificado
    else:
        project = Proyecto.query.filter_by(project_id=id_project).filter_by(user_id=id_user).first()
        data.append({
            "id": project.id,
            "name": project.name,
            "description": project.description,
            "status": project.status,
            "created_at": project.created_at,
            "updated_at": project.updated_at,
            "deleted_at": project.deleted_at
        })
    return {"status": "200", "message": "Proyectos obtenidos correctamente", "projects": data}


def get_project_user(data):
    id_project = data['id_project']
    id_user = data['id_user']
    project = Proyecto.query.filter_by(project_id=id_project).filter_by(user_id=id_user).first()
    data = {
        "id": project.project_id,
        "name": project.project_name,
        "description": project.description,
        "status": project.status,
        "created_at": project.created_at,
        "updated_at": project.updated_at
    }
    return {"status": "200", "message": "Proyecto obtenido correctamente", "project": data}


def get_files_project(data):
    id_project = data['id_project']
    id_user = data['id_user']
    project = Proyecto.query.filter_by(project_id=id_project).filter_by(user_id=id_user).first()
    if project is None:
        return {"status": "500", "message": "No se encontró el proyecto"}
    id_proyecto = project.project_id
    archivos = Archivo.query.filter_by(project_id=id_proyecto).all()
    if archivos is None:
        return {"status": "500", "message": "No se encontraron archivos"}
    files = []
    for archivo in archivos:
        files.append({
            "id": archivo.file_id,
            "name": archivo.filename,
            "original_name": archivo.original_filename,
            "created_at": archivo.created_at,
            "updated_at": archivo.updated_at,
        })
    return {"status": "200", "message": "Archivos obtenidos correctamente", "files": files}


def get_meta_data_file_project(data):
    id_project = data['id_project']
    id_user = data['id_user']
    id_file_ = data['id_file']
    project = Proyecto.query.filter_by(project_id=id_project).filter_by(user_id=id_user).first()

    """dir = project.dir
    for file in os.listdir(dir):
        sep = file.split("-")
        id_proyecto = sep[0]
        id_file = sep[1]
        # quitar los primeros dos eleeentos del array
        sep.pop(0)
        sep.pop(0)
        # unir el array en un string
        name = "-".join(sep)
        if id_file == id_file_:
            file = dir + "/" + file
            with open(file, 'r') as f:
                data = {'name': name, 'size': os.path.getsize(file), 'type': 'file'}
            break"""
    return {"status": "200", "message": "Metadatos obtenidos correctamente", "meta_data": "data"}


def delete_project(data):
    id_project = data['id_project']
    id_user = data['id_user']
    project = Proyecto.query.filter_by(project_id=id_project).filter_by(user_id=id_user).first()
    files = Archivo.query.filter_by(project_id=id_project).all()
    for file in files:
        #eliminar el archivo en el directorio
        directorio = current_app.config['UPLOADS']+"/"+file.filename
        os.remove(directorio)
        db.session.delete(file)
        db.session.commit()
    # elimminar el proyecto
    db.session.delete(project)
    db.session.commit()

    return {"status": "200", "message": "Proyecto eliminado correctamente"}


def get_file_text_project(data, fun=False, p_inicio=None, p_final=None):
    id_project = data['id_project']
    id_user = data['id_user']
    id_file = data['id_file']  # Ahora esto es un array
    project = Proyecto.query.filter_by(project_id=id_project).filter_by(user_id=id_user).first()
    if project is None:
        return {"status": "500", "message": "No se encontró el proyecto"}
    else:
        files = Archivo.query.filter(Archivo.file_id.in_(id_file)).filter_by(project_id=id_project).all()
        if not files:
            return {"status": "500", "message": "No se encontraron los archivos"}
        else:
            result = []
            combined_text = []
            for file in files:
                extension = file.filename.split(".")[-1]
                text = ""
                if extension == 'pdf':
                    text = read_pdf("uploads/" + file.filename, p_inicio, p_final)
                elif extension == 'docx':
                    text = read_docs("uploads/" + file.filename)
                combined_text.append(text)
            result.append({"status": "200", "message": "Archivo obtenido correctamente", "text": combined_text, 'file': file,
                           'name': file.filename})

            if fun:
                return combined_text  # Aquí quizá querrás cambiar el comportamiento también

            return result
    return {"status": "200", "message": "Archivo no encontrado"}


def quiz_batchSec(par: list, cross: False, sections: 64):
    questions = []
    answers = []
    predScore = []
    start_id = 0
    end_id = sections
    words = par.split()
    scale = math.trunc(len(words) / sections)
    if (cross):
        for i in range(scale):
            section_selected = ' '.join(words[start_id:end_id])
            result = generate_answer(str(section_selected))
            questions.append(result['question'])
            answers.append(result['answer'])
            predScore.append(result['score'])
            if (i < scale):
                start_id = start_id + int(sections / 2)
                end_id = start_id + sections
            if (i == scale):
                start_id = start_id + sections
                end_id = len(words) - (scale * i)
    if (not cross):
        for i in range(scale):
            section_selected = ' '.join(words[start_id:end_id])
            result = generate_answer(str(section_selected))
            questions.append(result['question'])
            answers.append(result['answer'])
            predScore.append(result['score'])
            start_id = end_id
            if (i < scale):
                end_id = start_id + sections
            if (i == scale):
                end_id = len(words) - 1
    return questions, answers, predScore


"""
def quiz_batchSec(par: str, cross: False, section_size: int = 64, batch_size: int = 1):
    questions = []
    answers = []
    predScore = []

    # Dividir el texto en secciones
    words = par.split()
    num_sections = math.ceil(len(words) / section_size)
    sections = [' '.join(words[i * section_size: (i + 1) * section_size]) for i in range(num_sections)]

    # Procesar en lotes (en este caso, una sección a la vez)
    for i in range(0, len(sections), batch_size):
        batch_sections = sections[i: i + batch_size]

        # Concatenar secciones en el lote (solo una sección en este caso)
        batch_text = ' '.join(batch_sections)

        # Generar preguntas y respuestas para el lote
        result = generate_answer(batch_text)
        questions.append(result['question'])
        answers.append(result['answer'])
        predScore.append(result['score'])

    return questions, answers, predScore
"""

def document_to_quiz(q: list, a: list) -> object:
    document = Document()

    document.add_heading('Exámen Generado', 0)

    p = document.add_paragraph('Tu examen generado es el siguiente:').bold = True
    document.add_paragraph('Tus preguntas son', style='List Bullet')
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pregunta'
    hdr_cells[1].text = 'Respuesta'

    for i in range(len(q)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(q[i])
        row_cells[1].text = str(a[i])

    dir = "temp"
    t = time.localtime()
    nhash = hash(str(t.tm_year) + str(t.tm_mon) + str(t.tm_mday) + str(t.tm_hour) + str(t.tm_min) + str(t.tm_sec))
    in_doc = f'{dir}/demo-{nhash}.docx'
    out_doc = f'{dir}/demo-{nhash}.pdf'
    name = f'demo-{nhash}.pdf'

    document.save(in_doc)
    doc = aw.Document(in_doc)
    doc.save(out_doc)
    return dir, name


def download_file_project(id_project, id_file, id_user):
    # id_project = data['id_project']
    # id_user = data['id_user']
    # id_file = data['id_file']
    project = Proyecto.query.filter_by(project_id=id_project).filter_by(user_id=id_user).first()
    if not project:
        return {"status": "400", "message": "Proyecto no encontrado"}
    else:
        file = Archivo.query.filter_by(file_id=id_file).first()
        if not file:
            return {"status": "400", "message": "Archivo no encontrado"}
        else:
            return send_from_directory("uploads", file.filename, as_attachment=True)
    # return {"status": "200", "message": "Archivo no encontrado"}


def download_file_project_tmp(name, id_user):
    f = name.split("/")
    namef = f[1]
    dir = f[0]

    return send_from_directory(dir, namef, as_attachment=True)


def descargar_archivo(dir, nombre_archivo):
    # Obtener la ruta completa del archivo
    ruta_archivo = os.path.join(dir, nombre_archivo)

    # Verificar si el archivo existe
    if not os.path.exists(ruta_archivo):
        # Si no existe, devolver un mensaje de error
        return "El archivo no existe", 404

    # Si el archivo existe, leerlo y devolverlo en la respuesta
    with open(ruta_archivo, 'rb') as f:
        contenido = f.read()
    return send_file(
        BytesIO(contenido),
        mimetype='application/octet-stream'
    )


def read_file_project(data):
    id_project = data['id_project']
    id_user = data['id_user']
    id_file = data['id_file']
    project = Projects.query.filter_by(id=id_project).filter_by(id_user=id_user).first()
    dir = project.dir
    f = []
    if not os.path.exists(dir):
        os.makedirs(dir)
    files = os.listdir(dir)
    for file in files:
        sep = file.split("-")
        id_proyecto = sep[0]
        id_file_ = sep[1]
        # quitar los primeros dos eleeentos del array
        sep.pop(0)
        sep.pop(0)
        # unir el array en un string
        name = "-".join(sep)
        if id_file == id_file_:
            # retornar el arcxhivo para descargar
            return send_from_directory(dir, file, as_attachment=False)
    return {"status": "200", "message": "Archivo no encontrado"}


"""
def createPDF():
    pdf = FPDF()

    # Añadimos una página al PDF
    pdf.add_page()

    # Definimos la fuente y el tamaño para el título
    pdf.set_font("Arial", size=24, style="B")
    pdf.cell(0, 20, txt="Examen generado", ln=1, align="C")

    # Definimos la fuente y el tamaño para las preguntas y respuestas
    pdf.set_font("Arial", size=16)
    pdf.set_text_color(0, 0, 255)

    # Definimos el color de fondo para las preguntas y respuestas
    pdf.set_fill_color(220, 220, 220)

    # Creamos una lista de preguntas y respuestas
    preguntas_respuestas = [
        {
            "pregunta": "¿Cuál es la capital de Francia?",
            "respuesta": "París"
        },
        {
            "pregunta": "¿En qué año se fundó Apple?",
            "respuesta": "1976"
        },
        {
            "pregunta": "¿Cuál es el río más largo del mundo?",
            "respuesta": "El río Amazonas"
        }
    ]

    # Recorremos la lista de preguntas y respuestas y las escribimos en el PDF
    for i, pregunta_respuesta in enumerate(preguntas_respuestas):
        pregunta = pregunta_respuesta["pregunta"]
        respuesta = pregunta_respuesta["respuesta"]
        pdf.cell(0, 10, txt=f"{i + 1}. {pregunta}\n", fill=True)
        pdf.cell(0, 10, txt=f"Respuesta: {respuesta}", ln=1)
    # Guardamos el PDF
    # gneerar un hash para el nombre del archivo
"""
def new_create_PDF(questions, answers):
    # Abre el documento y obtén la página que deseas modificar.
    doc = fitz.open("assets/file_base.pdf")
    page = doc[0]

    # Inserta el título en la parte superior de la página.
    title = "Examen generado"
    title_location = fitz.Point(50, 72)  # ajusta las coordenadas como sea necesario
    title_color = (0, 0.4, 0.9)
    page.insert_text(title_location, title, fontname="helv", fontsize=20, color=title_color, overlay=True)

    # Define los puntos de inicio y fin de la línea y dibuja la línea en la página.
    start = fitz.Point(50, 80)  # ajusta las coordenadas como sea necesario
    end = fitz.Point(550, 80)  # ajusta las coordenadas como sea necesario
    line = page.add_line_annot(start, end)
    line.set_border(width=1.2)
    line.set_colors(stroke=(0, 0.6, 0.9))  # color azul claro
    line.update()

    # Supón que tienes las preguntas y respuestas en dos listas, `preguntas` y `respuestas`.
    preguntas = questions#['¿Pregunta 1?', '¿Pregunta 2?', '¿Pregunta 3?']
    respuestas = answers#['Respuesta 1', 'Respuesta 2', 'Respuesta 3']

    y_position = 120  # posición inicial de la y para la primera pregunta

    for i in range(len(preguntas)):
        question = f"{i + 1}. {preguntas[i]}"
        answer = f"Respuesta: {respuestas[i]}"

        # Inserta la pregunta
        question_location = fitz.Point(50, y_position)
        page.insert_text(question_location, question, fontname="helv", fontsize=11)

        # Inserta la respuesta
        y_position += 20  # ajusta según el espacio que quieras entre preguntas y respuestas
        answer_location = fitz.Point(50, y_position)
        page.insert_text(answer_location, answer, fontname="helv", fontsize=11)

        y_position += 40  # ajusta según el espacio que quieras entre las preguntas

    # Guarda el documento.
    doc.save("temp/text.pdf")
    """dir = "temp"
    t = time.localtime()
    nhash = hash(str(t.tm_year) + str(t.tm_mon) + str(t.tm_mday) + str(t.tm_hour) + str(t.tm_min) + str(t.tm_sec))
    in_doc = f'{dir}/demo-{nhash}.docx'
    out_doc = f'{dir}/demo-{nhash}-fritz.pdf'
    name = f'demo-{nhash}-fritz.pdf'

    #document.save(in_doc)
    #doc = aw.Document(in_doc)
    doc.save(out_doc)
    return dir, name"""
    return "temp", "demo.pdf"



def topQuestions(question: list, answer: list, score: list, top: int):
    data = {'question': question, 'answer': answer, 'score': score}
    df = pd.DataFrame(data)
    df = df.sort_values(by='score', ascending=False)
    df = df.head(top)
    questions = df['question'].tolist()
    answers = df['answer'].tolist()
    score = df['score'].tolist()
    return questions, answers


def createQuiz(data, id, id_file, id_proyecto, cantidad_preguntas=10, reemplazar=False, p_inicio=None, p_final=None):
    print(data, id, id_file, id_proyecto, cantidad_preguntas)
    texto = get_file_text_project(data, fun=True, p_inicio=p_inicio, p_final=p_final)
    texto_unido = ' '.join([' '.join(pagina) for pagina in texto])
    q, a, score = quiz_batchSec(texto_unido, False, 128) #128
    topQ, topA = topQuestions(q, a, score, cantidad_preguntas)
    try:
        parafraser_question(topQ)
    except:
        pass
    # name, dir = document_to_quiz(topQ, topA)
    saveFileInDb(id_proyecto, id_file, [topQ, topA], reemplazar, data)
    # createPDF()
    # return {"status": "400", "message": "Error al crear el quiz"}
    return {"status": "200", "message": "Quiz creado correctamente", "quiz questions": topQ,
            "quiz answers": topA}
def parafraser_question(questions):
    print("Preguntas: ", questions)
    n = int(len(questions) * 0.2)

    # Seleccionar índices del 20% de los elementos de forma aleatoria sin repetición
    indices_seleccionados = random.sample(range(len(questions)), n)

    # Parafrasear y reemplazar en el array original
    for index in indices_seleccionados:
        pregunta_original = questions[index]
        lista_parafraseado = generate(pregunta_original)
        diccionario_parafraseado = random.choice(lista_parafraseado)  # Elegir un diccionario al azar
        pregunta_parafraseada = diccionario_parafraseado["respuesta"]
        questions[index] = pregunta_parafraseada
        print("Pregunta inicial: ", pregunta_original, " Pregunta parafraseada:", pregunta_parafraseada)

    # Mostrar el array modificado
    print("Preguntas con parafraseado: ", questions)
    return questions

def saveFileInDb(id_proyecto, id_files, data, reemplazar, json_data):
    id = json_data['id_user']

    cantidad = json_data['cant_questions']
    reescribir = json_data['reescribir']
    id_cuestionario_antiguo = json_data['id_cuestionario_antiguo']
    #names files
    files = Archivo.query.filter(Archivo.file_id.in_(id_files)).all()
    names = [file.original_filename for file in files]


    if reemplazar:
        #rremplazar el test existente
        test = Test.query.filter_by(test_id=id_cuestionario_antiguo).first()
        if test:
            test.test_data = json.dumps({"questions": data[0], "answers": data[1]})
            test.updated_at = datetime.now()
            test.names_files = names
            db.session.commit()
    # Crear un nuevo objeto Test (cuestionario)
    else:
        new_test = Test(
            proyecto_id=id_proyecto,
            test_data=json.dumps({"questions": data[0], "answers": data[1]}),
            names_files=names,
            is_multi_file=len(id_files) > 1  # establecer is_multi_file según el número de archivos
        )

        # Añadir los archivos relacionados
        for file_id in id_files:
            file = Archivo.query.get(file_id)
            if file:  # Asegúrate de que el archivo exista antes de agregarlo a la relación
                new_test.archivos.append(file)

        # Agregar el nuevo cuestionario a la sesión de la base de datos
        db.session.add(new_test)

        # Confirmar los cambios en la base de datos
        db.session.commit()


def generate_pdf(data):
    questions = data.get('q', [])
    answers = data.get('a', [])
    d1, n1 = new_create_PDF(questions, answers)
    directorio, name = document_to_quiz(questions, answers)

    print(d1, n1)
    return send_from_directory(directorio, name, as_attachment=True)


def borrar_archivos(directorio):
    for filename in os.listdir(directorio):
        file_path = os.path.join(directorio, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                os.rmdir(file_path)
        except Exception as e:
            print('Error al borrar %s. Razón: %s' % (file_path, e))


def get_all_quiz(data):
    id_proyecto = data['id_project']
    # id_user = data['id_user']
    # id_file = data['id_file']

    project = Test.query.filter_by(proyecto_id=id_proyecto).all()
    if project:
        data_n = []
        for p in project:
            try:
                data_n.append({'cuestionario': json.loads(p.test_data), 'created_at': p.created_at.strftime("%d/%m/%Y %H:%M:%S"), 'name_files': p.names_files, 'id_cuestionario': p.test_id})
            except json.JSONDecodeError as e:
                print(f"Error parsing JSON: {e}")
        return {"status": "200", "message": "Quiz encontrado", "quiz": data_n}
    return {"status": "200", "message": "Quiz no encontrado"}


def delete_file(data):
    id_file = data['id_file']
    id_user = data['id_user']
    id_proyecto = data['id_project']

    file = Archivo.query.filter_by(file_id=id_file).first()
    if file:
        try:
            os.remove(os.path.join(current_app.config['UPLOADS'], file.filename))
            db.session.delete(file)
            db.session.commit()
            return {"status": "200", "message": "Archivo eliminado correctamente"}
        except Exception as e:
            print(f"Error al eliminar el archivo: {e}")
            return {"status": "400", "message": "Error al eliminar el archivo"}
    return {"status": "400", "message": "Archivo no encontrado"}


def get_quiz_by_id(id):
    quiz = Test.query.get(id)
    if quiz:
        try:
            data = json.loads(quiz.test_data)
            return {"status": "200", "message": "Quiz encontrado", "quiz": data}
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON: {e}")
    return {"status": "400", "message": "Quiz no encontrado"}


class QuestionAnswerTemplate(Flowable):
    def __init__(self, questions, answers):
        Flowable.__init__(self)
        self.questions = questions
        self.answers = answers

    def draw(self):
        c = self.canv
        for i, (question, answer) in enumerate(zip(self.questions, self.answers)):
            c.drawString(100, 800 - i * 100, question)
            c.drawString(300, 800 - i * 100, answer)
